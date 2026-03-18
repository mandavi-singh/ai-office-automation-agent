"""
OCR tools for extracting text from images and saving the output.
"""
import os


SUPPORTED_IMAGE_EXTENSIONS = {
    ".jpg",
    ".jpeg",
    ".png",
    ".bmp",
    ".tiff",
    ".tif",
    ".gif",
    ".webp",
}


def _workspace_output_dir() -> str:
    output_dir = os.path.join(os.getcwd(), "output")
    os.makedirs(output_dir, exist_ok=True)
    return output_dir


def _normalize_output_file(path: str, default_name: str, extension: str) -> str:
    if not path:
        return os.path.join(_workspace_output_dir(), default_name)

    abs_path = os.path.abspath(path)
    if os.path.isdir(abs_path):
        return os.path.join(abs_path, default_name)
    if not abs_path.lower().endswith(extension.lower()):
        return abs_path + extension
    return abs_path


def _configure_tesseract(pytesseract) -> None:
    """Point pytesseract at common Windows install locations when available."""
    tesseract_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        r"C:\Users\singh\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
    ]
    for path in tesseract_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            return


def ocr_image_to_text(image_path: str, language: str = "eng", save_to: str = "") -> dict:
    """Extract text from an image using Tesseract OCR."""
    try:
        if not os.path.exists(image_path):
            return {"success": False, "message": f"File not found: {image_path}", "text": ""}

        ext = os.path.splitext(image_path)[1].lower()
        if ext not in SUPPORTED_IMAGE_EXTENSIONS:
            return {
                "success": False,
                "message": f"Unsupported format: {ext}. Use JPG, PNG, BMP, TIFF, GIF, or WEBP.",
                "text": "",
            }

        try:
            import pytesseract
            from PIL import Image, ImageOps
        except ImportError:
            return {
                "success": False,
                "message": (
                    "pytesseract not installed.\n"
                    "Run: pip install pytesseract pillow\n"
                    "Also install Tesseract OCR from: https://github.com/UB-Mannheim/tesseract/wiki"
                ),
                "text": "",
            }

        _configure_tesseract(pytesseract)

        img = Image.open(image_path)
        img = ImageOps.exif_transpose(img).convert("L")
        text = pytesseract.image_to_string(img, lang=language).strip()

        if not text:
            return {
                "success": False,
                "message": "No text found in image. Make sure the image is clear and readable.",
                "text": "",
            }

        if save_to:
            save_to = _normalize_output_file(
                save_to,
                f"{os.path.splitext(os.path.basename(image_path))[0]}_ocr.txt",
                ".txt",
            )
            os.makedirs(os.path.dirname(os.path.abspath(save_to)), exist_ok=True)
            with open(save_to, "w", encoding="utf-8") as handle:
                handle.write(text)
            save_msg = f"\nSaved to: {save_to}"
        else:
            save_msg = ""

        return {
            "success": True,
            "text": text,
            "char_count": len(text),
            "message": (
                f"OCR complete. Extracted {len(text)} characters from "
                f"{os.path.basename(image_path)}{save_msg}"
            ),
        }
    except Exception as exc:
        if "tesseract" in str(exc).lower():
            return {
                "success": False,
                "message": (
                    "Tesseract OCR engine not found.\n"
                    "Install it from: https://github.com/UB-Mannheim/tesseract/wiki"
                ),
                "text": "",
            }
        return {"success": False, "message": f"OCR error: {exc}", "text": ""}


def ocr_image_to_word(image_path: str, output_docx: str, language: str = "eng") -> dict:
    """Extract text from an image and save it to a Word file."""
    try:
        result = ocr_image_to_text(image_path=image_path, language=language)
        if not result.get("success"):
            return result

        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(f"OCR Output - {os.path.basename(image_path)}", level=1)

        for line in result["text"].splitlines():
            if not line.strip():
                continue
            paragraph = doc.add_paragraph(line)
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(12)

        output_docx = _normalize_output_file(
            output_docx,
            f"{os.path.splitext(os.path.basename(image_path))[0]}_ocr.docx",
            ".docx",
        )
        os.makedirs(os.path.dirname(os.path.abspath(output_docx)), exist_ok=True)
        doc.save(output_docx)

        return {
            "success": True,
            "text": result["text"],
            "char_count": result.get("char_count", len(result["text"])),
            "message": f"OCR text saved to Word file: {output_docx}",
            "output_docx": output_docx,
        }
    except Exception as exc:
        return {"success": False, "message": f"Error saving Word file: {exc}", "text": ""}
