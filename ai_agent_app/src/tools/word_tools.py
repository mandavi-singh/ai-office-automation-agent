"""
MS Word automation tools.
Uses python-docx for document operations.
"""
import os

from docx import Document
from docx.shared import Pt, RGBColor


def _hex_to_rgb(hex_color: str):
    """Convert a hex string like FF0000 to an RGBColor."""
    cleaned = hex_color.lstrip("#")
    return RGBColor(int(cleaned[0:2], 16), int(cleaned[2:4], 16), int(cleaned[4:6], 16))


def _open_created_word_file(filename: str) -> str:
    """Open a saved Word document with the default Windows app when possible."""
    try:
        absolute = os.path.abspath(filename)
        os.startfile(absolute)
        return f" Opened in Word: {absolute}"
    except Exception:
        return ""


def word_create_file(filename: str, title: str, title_font_size: int = 24,
                     content: str = "", font_color: str = "") -> str:
    """Create a new Word document with a title and optional body content."""
    try:
        os.makedirs(os.path.dirname(os.path.abspath(filename)), exist_ok=True)
        doc = Document()

        title_para = doc.add_heading(title, level=0)
        title_run = title_para.runs[0] if title_para.runs else title_para.add_run(title)
        title_run.font.size = Pt(title_font_size)
        title_run.bold = True
        if font_color:
            title_run.font.color.rgb = _hex_to_rgb(font_color)

        if content:
            for paragraph in content.split("\n"):
                if not paragraph.strip():
                    continue
                para = doc.add_paragraph()
                run = para.add_run(paragraph)
                if font_color:
                    run.font.color.rgb = _hex_to_rgb(font_color)

        doc.save(filename)
        opened_message = _open_created_word_file(filename)
        return f"✅ Word file created: {filename}.{opened_message}"
    except Exception as e:
        return f"❌ Error creating Word file: {e}"


def word_format_text(filename: str, search_text: str, bold: bool = None,
                     italic: bool = None, font_color: str = "", font_size: int = None) -> str:
    """Find matching text in a Word file and apply formatting."""
    try:
        doc = Document(filename)
        count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if search_text.lower() in run.text.lower():
                    if bold is not None:
                        run.bold = bold
                    if italic is not None:
                        run.italic = italic
                    if font_color:
                        run.font.color.rgb = _hex_to_rgb(font_color)
                    if font_size:
                        run.font.size = Pt(font_size)
                    count += 1
        doc.save(filename)
        if count:
            return f"✅ Formatted {count} run(s) containing '{search_text}' in {filename}"
        return f"⚠️ Text '{search_text}' not found in {filename}"
    except Exception as e:
        return f"❌ Error formatting Word file: {e}"


def word_add_content(filename: str, text: str, style: str = "Normal",
                     bold: bool = False, italic: bool = False,
                     font_color: str = "", font_size: int = None) -> str:
    """Add a paragraph or heading to an existing Word document."""
    try:
        doc = Document(filename)
        if style.startswith("Heading"):
            level = int(style.split()[-1]) if style[-1].isdigit() else 1
            para = doc.add_heading(text, level=level)
            if para.runs:
                run = para.runs[0]
                run.bold = bold or run.bold
                run.italic = italic
                if font_color:
                    run.font.color.rgb = _hex_to_rgb(font_color)
                if font_size:
                    run.font.size = Pt(font_size)
        else:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = bold
            run.italic = italic
            if font_color:
                run.font.color.rgb = _hex_to_rgb(font_color)
            if font_size:
                run.font.size = Pt(font_size)

        doc.save(filename)
        return f"✅ Added paragraph to {filename}"
    except Exception as e:
        return f"❌ Error adding content to Word file: {e}"
